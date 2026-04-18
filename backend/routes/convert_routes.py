import os
import uuid
import io
import struct
import zlib
from flask import Blueprint, request, jsonify, current_app, send_from_directory
from flask_jwt_extended import jwt_required, get_jwt_identity
from werkzeug.utils import secure_filename
from PIL import Image
from moviepy import VideoFileClip
from pydub import AudioSegment
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from pdf2image import convert_from_path
import pillow_heif
import vtracer
import re

pillow_heif.register_heif_opener()

from models import Job
from extensions import db

convert_bp = Blueprint('convert', __name__)

ALLOWED_EXTENSIONS = {
    'png', 'jpg', 'jpeg', 'webp', 'bmp', 'gif', 'tiff', 'tif',
    'svg', 'heic', 'jfif', 'pdf', 'docx', 'mp4', 'mp3', 'mov',
    'avi', 'ogg', 'wav', 'm4a', 'flac', 'mkv', 'wmv', 'webm'
}

IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'webp', 'bmp', 'gif', 'tiff', 'tif', 'heic', 'jfif'}
AUDIO_EXTENSIONS = {'mp3', 'ogg', 'wav', 'flac', 'm4a'}
VIDEO_EXTENSIONS = {'mp4', 'mov', 'avi', 'mkv', 'wmv', 'webm'}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def get_upload_folder():
    return current_app.config['UPLOAD_FOLDER']


# ─────────────────────────────────────────────
# UPLOAD
# ─────────────────────────────────────────────

@convert_bp.route('/upload', methods=['POST'])
@jwt_required(optional=True)
def upload_file():
    if 'file' not in request.files:
        return jsonify({'message': 'No file part'}), 400

    file = request.files['file']
    target_format = request.form.get('target_format', '').lower()

    if file.filename == '':
        return jsonify({'message': 'No selected file'}), 400

    if not target_format or target_format not in ALLOWED_EXTENSIONS:
        return jsonify({'message': 'Invalid target format'}), 400

    if file and allowed_file(file.filename):
        user_identity = get_jwt_identity()
        original_filename = secure_filename(file.filename)
        unique_id = str(uuid.uuid4())
        ext = original_filename.rsplit('.', 1)[1].lower()
        save_name = f"{unique_id}.{ext}"
        filepath = os.path.join(get_upload_folder(), save_name)
        file.save(filepath)

        new_job = Job(
            user_id=user_identity,
            original_filename=save_name,
            target_format=target_format,
            status='pending'
        )
        db.session.add(new_job)
        db.session.commit()
        return jsonify({'message': 'File uploaded', 'job_id': new_job.id}), 201

    return jsonify({'message': 'File type not allowed'}), 400


# ─────────────────────────────────────────────
# PROCESS
# ─────────────────────────────────────────────

@convert_bp.route('/process/<int:job_id>', methods=['POST'])
def process_job(job_id):
    job = db.get_or_404(Job, job_id)
    if job.status != 'pending':
        return jsonify({'message': 'Job already processed'}), 400

    job.status = 'processing'
    db.session.commit()

    try:
        target_ext = job.target_format
        upload_folder = get_upload_folder()

        input_path = os.path.join(upload_folder, job.original_filename)
        original_ext = job.original_filename.rsplit('.', 1)[1].lower()
        output_filename = f"{job.original_filename.rsplit('.', 1)[0]}.{target_ext}"
        output_path = os.path.join(upload_folder, output_filename)

        # ── Compression jobs ───────────────────────────────────────────────
        compress_level = request.json.get('quality', None) if request.is_json else None

        if target_ext == 'compress':
            output_filename, output_path = _compress_file(
                input_path, original_ext, upload_folder, compress_level
            )
        # ── PDF → image ────────────────────────────────────────────────────
        elif target_ext in IMAGE_EXTENSIONS and original_ext == 'pdf':
            _pdf_to_image(input_path, output_path, target_ext)
        # ── SVG → raster ───────────────────────────────────────────────────
        elif original_ext == 'svg' and target_ext in IMAGE_EXTENSIONS:
            _svg_to_image(input_path, output_path, target_ext)
        # ── SVG → PDF ──────────────────────────────────────────────────────
        elif original_ext == 'svg' and target_ext == 'pdf':
            try:
                import cairosvg
                cairosvg.svg2pdf(url=input_path, write_to=output_path)
            except ImportError:
                raise ValueError("SVG to PDF conversion requires cairosvg and Cairo library. Please install Cairo on your system.")
        # ── image → SVG ────────────────────────────────────────────────────
        elif target_ext == 'svg' and original_ext in IMAGE_EXTENSIONS:
            vtracer.convert_image_to_svg_py(input_path, output_path)
        # ── image → image ──────────────────────────────────────────────────
        elif target_ext in IMAGE_EXTENSIONS and original_ext in IMAGE_EXTENSIONS:
            _image_to_image(input_path, output_path, target_ext)
        # ── audio ──────────────────────────────────────────────────────────
        elif target_ext in AUDIO_EXTENSIONS:
            _convert_audio(input_path, output_path, original_ext, target_ext)
        # ── video ──────────────────────────────────────────────────────────
        elif target_ext in VIDEO_EXTENSIONS:
            _convert_video(input_path, output_path, original_ext, target_ext)
        # ── PDF ────────────────────────────────────────────────────────────
        elif target_ext == 'pdf':
            _to_pdf(input_path, output_path, original_ext)
        # ── DOCX ───────────────────────────────────────────────────────────
        elif target_ext == 'docx':
            _to_docx(input_path, output_path, original_ext)
        else:
            raise ValueError(f"Conversion from {original_ext} to {target_ext} not supported")

        # ── Optional compression after conversion ────────────────────────────
        if compress_level is not None and target_ext != 'compress':
            compressed_filename, compressed_path = _compress_file(
                output_path, target_ext, upload_folder, compress_level
            )
            if compressed_path != output_path and os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except OSError:
                    pass
            output_filename = compressed_filename
            output_path = compressed_path

        job.status = 'done'
        job.result_filename = output_filename
        db.session.commit()
        return jsonify({'message': 'Job complete', 'job_id': job.id, 'status': 'done'}), 200

    except Exception as e:
        import traceback
        traceback.print_exc()
        job.status = 'failed'
        db.session.commit()
        return jsonify({'message': f'Conversion failed: {str(e)}'}), 500


# ─────────────────────────────────────────────
# COMPRESS endpoint (separate route for flexibility)
# ─────────────────────────────────────────────

@convert_bp.route('/compress/<int:job_id>', methods=['POST'])
def compress_job(job_id):
    """Dedicated compression endpoint. Accepts optional JSON body: {"quality": 60}"""
    job = db.get_or_404(Job, job_id)
    if job.status != 'pending':
        return jsonify({'message': 'Job already processed'}), 400

    job.status = 'processing'
    db.session.commit()

    try:
        data = request.get_json(silent=True) or {}
        quality = int(data.get('quality', 70))
        upload_folder = get_upload_folder()
        input_path = os.path.join(upload_folder, job.original_filename)
        original_ext = job.original_filename.rsplit('.', 1)[1].lower()

        output_filename, output_path = _compress_file(
            input_path, original_ext, upload_folder, quality
        )

        job.status = 'done'
        job.result_filename = output_filename
        db.session.commit()
        return jsonify({'message': 'Compression complete', 'job_id': job.id, 'status': 'done'}), 200

    except Exception as e:
        import traceback
        traceback.print_exc()
        job.status = 'failed'
        db.session.commit()
        return jsonify({'message': f'Compression failed: {str(e)}'}), 500


# ─────────────────────────────────────────────
# STATUS / DOWNLOAD
# ─────────────────────────────────────────────

@convert_bp.route('/status/<int:job_id>', methods=['GET'])
def get_status(job_id):
    job = db.get_or_404(Job, job_id)
    return jsonify({
        'job_id': job.id,
        'status': job.status,
        'original_filename': job.original_filename,
        'target_format': job.target_format
    }), 200


@convert_bp.route('/download/<int:job_id>', methods=['GET'])
def download_file(job_id):
    job = db.get_or_404(Job, job_id)
    if job.status == 'done' and job.result_filename:
        return send_from_directory(
            os.path.abspath(get_upload_folder()),
            job.result_filename,
            as_attachment=True
        )
    return jsonify({'message': 'File not ready or job failed'}), 400


# ═════════════════════════════════════════════
# PRIVATE HELPERS
# ═════════════════════════════════════════════

# ── Image helpers ─────────────────────────────

def _image_to_image(input_path, output_path, target_ext):
    with Image.open(input_path) as img:
        # Convert palette/transparency for JPEG family
        if target_ext in ('jpg', 'jpeg', 'jfif') and img.mode in ('RGBA', 'P', 'LA'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            if img.mode in ('RGBA', 'LA'):
                background.paste(img, mask=img.split()[-1])
            img = background
        elif target_ext == 'png' and img.mode == 'P':
            img = img.convert('RGBA')

        fmt_map = {'jpg': 'JPEG', 'jpeg': 'JPEG', 'jfif': 'JPEG', 'tif': 'TIFF'}
        fmt = fmt_map.get(target_ext, target_ext.upper())
        img.save(output_path, format=fmt)


def _pdf_to_image(input_path, output_path, target_ext):
    images = convert_from_path(input_path, dpi=150)
    if not images:
        raise ValueError("Could not extract pages from PDF")
    img = images[0]
    if target_ext in ('jpg', 'jpeg') and img.mode == 'RGBA':
        bg = Image.new('RGB', img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[3])
        img = bg
    fmt_map = {'jpg': 'JPEG', 'jpeg': 'JPEG', 'jfif': 'JPEG', 'tif': 'TIFF'}
    fmt = fmt_map.get(target_ext, target_ext.upper())
    img.save(output_path, format=fmt)


def _svg_to_image(input_path, output_path, target_ext):
    try:
        import cairosvg
    except ImportError:
        raise ValueError("SVG conversion requires cairosvg and Cairo library. Please install Cairo on your system.")
    # Render SVG to PNG first, then convert
    png_bytes = cairosvg.svg2png(url=input_path)
    img = Image.open(io.BytesIO(png_bytes))
    if target_ext in ('jpg', 'jpeg') and img.mode == 'RGBA':
        bg = Image.new('RGB', img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[3])
        img = bg
    fmt_map = {'jpg': 'JPEG', 'jpeg': 'JPEG', 'jfif': 'JPEG', 'tif': 'TIFF'}
    fmt = fmt_map.get(target_ext, target_ext.upper())
    img.save(output_path, format=fmt)


# ── Audio helpers ─────────────────────────────

def _convert_audio(input_path, output_path, original_ext, target_ext):
    if original_ext in AUDIO_EXTENSIONS:
        audio = AudioSegment.from_file(input_path)
    elif original_ext in VIDEO_EXTENSIONS:
        video = VideoFileClip(input_path)
        tmp_wav = input_path + '_tmp.wav'
        video.audio.write_audiofile(tmp_wav, logger=None)
        audio = AudioSegment.from_wav(tmp_wav)
        os.remove(tmp_wav)
        video.close()
    else:
        raise ValueError(f"Cannot extract audio from {original_ext}")

    fmt_map = {'m4a': 'ipod'}
    fmt = fmt_map.get(target_ext, target_ext)
    audio.export(output_path, format=fmt)


# ── Video helpers ─────────────────────────────

def _convert_video(input_path, output_path, original_ext, target_ext):
    if original_ext not in VIDEO_EXTENSIONS:
        raise ValueError(f"Cannot convert {original_ext} to video")
    video = VideoFileClip(input_path)
    codec_map = {
        'mp4': 'libx264',
        'webm': 'libvpx',
        'avi': None,
        'mov': 'libx264',
        'mkv': 'libx264',
        'wmv': 'wmv2',
    }
    codec = codec_map.get(target_ext)
    kwargs = {'logger': None}
    if codec:
        kwargs['codec'] = codec
    video.write_videofile(output_path, **kwargs)
    video.close()


# ── PDF helpers ───────────────────────────────

def _to_pdf(input_path, output_path, original_ext):
    if original_ext == 'docx':
        _docx_to_pdf(input_path, output_path)
    elif original_ext in IMAGE_EXTENSIONS:
        img = Image.open(input_path).convert('RGB')
        img.save(output_path, 'PDF', resolution=100)
    elif original_ext == 'svg':
        try:
            import cairosvg
            cairosvg.svg2pdf(url=input_path, write_to=output_path)
        except ImportError:
            raise ValueError("SVG to PDF conversion requires cairosvg and Cairo library. Please install Cairo on your system.")
    else:
        raise ValueError(f"Cannot convert {original_ext} to PDF")


def _docx_to_pdf(input_path, output_path):
    """Convert DOCX to PDF preserving paragraphs, headings, bold/italic, tables."""
    doc = Document(input_path)
    story = []
    styles = getSampleStyleSheet()

    heading_style = ParagraphStyle('Heading1RL', parent=styles['Heading1'], fontSize=16, spaceAfter=12)
    heading2_style = ParagraphStyle('Heading2RL', parent=styles['Heading2'], fontSize=14, spaceAfter=10)
    normal_style = ParagraphStyle('NormalRL', parent=styles['Normal'], fontSize=11, spaceAfter=6, leading=14)
    bold_style = ParagraphStyle('BoldRL', parent=styles['Normal'], fontSize=11, spaceAfter=6, fontName='Helvetica-Bold')

    def _para_to_rl(para):
        # Detect heading style
        if para.style.name.startswith('Heading 1'):
            return Paragraph(_escape_xml(para.text), heading_style)
        if para.style.name.startswith('Heading 2'):
            return Paragraph(_escape_xml(para.text), heading2_style)

        # Build inline run markup
        parts = []
        for run in para.runs:
            text = _escape_xml(run.text)
            if not text:
                continue
            if run.bold and run.italic:
                parts.append(f'<b><i>{text}</i></b>')
            elif run.bold:
                parts.append(f'<b>{text}</b>')
            elif run.italic:
                parts.append(f'<i>{text}</i>')
            else:
                parts.append(text)
        return Paragraph(''.join(parts), normal_style)

    for element in doc.element.body:
        tag = element.tag.split('}')[-1]
        if tag == 'p':
            from docx.oxml.ns import qn
            from docx.text.paragraph import Paragraph as DocxPara
            para = DocxPara(element, doc)
            rl_para = _para_to_rl(para)
            story.append(rl_para)
            story.append(Spacer(1, 2))
        elif tag == 'tbl':
            from docx.table import Table as DocxTable
            table = DocxTable(element, doc)
            data = []
            for row in table.rows:
                data.append([_escape_xml(cell.text) for cell in row.cells])
            if data:
                rl_table = Table(data, hAlign='LEFT')
                rl_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6366f1')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')]),
                    ('PADDING', (0, 0), (-1, -1), 6),
                ]))
                story.append(rl_table)
                story.append(Spacer(1, 8))

    pdf = SimpleDocTemplate(output_path, pagesize=A4,
                            leftMargin=inch, rightMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    pdf.build(story)


def _escape_xml(text):
    if not text:
        return ''
    return (text.replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;'))


# ── DOCX helpers ──────────────────────────────

def _to_docx(input_path, output_path, original_ext):
    if original_ext == 'pdf':
        _pdf_to_docx(input_path, output_path)
    elif original_ext in IMAGE_EXTENSIONS:
        _image_to_docx(input_path, output_path)
    else:
        raise ValueError(f"Cannot convert {original_ext} to DOCX")


def _pdf_to_docx(input_path, output_path):
    """Extract PDF text with layout hints and write to DOCX."""
    reader = PdfReader(input_path)
    doc = Document()

    # Title style
    doc.styles['Normal'].font.size = Pt(11)

    for page_num, page in enumerate(reader.pages):
        if page_num > 0:
            doc.add_page_break()

        text = page.extract_text() or ''
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph('')
                continue

            para = doc.add_paragraph()
            # Heuristic: short ALL-CAPS lines or lines ending with no punctuation
            # that are significantly shorter than average → treat as heading
            if len(line) < 80 and line.isupper():
                para.style = doc.styles['Heading 1']
                para.add_run(line)
            elif len(line) < 60 and not line.endswith(('.', ',', ';', ':')):
                para.style = doc.styles['Heading 2']
                para.add_run(line)
            else:
                para.add_run(line)

    doc.save(output_path)


def _image_to_docx(input_path, output_path):
    doc = Document()
    doc.add_heading('Image', level=1)
    doc.add_picture(input_path, width=Inches(5.5))
    doc.save(output_path)


# ── Compression helpers ───────────────────────

def _compress_file(input_path, original_ext, upload_folder, quality=None):
    unique_id = str(uuid.uuid4())

    if original_ext in IMAGE_EXTENSIONS:
        return _compress_image(input_path, original_ext, upload_folder, unique_id, quality or 70)
    elif original_ext == 'pdf':
        return _compress_pdf(input_path, upload_folder, unique_id, quality or 70)
    elif original_ext in AUDIO_EXTENSIONS:
        return _compress_audio(input_path, original_ext, upload_folder, unique_id, quality or 128)
    elif original_ext in VIDEO_EXTENSIONS:
        return _compress_video(input_path, original_ext, upload_folder, unique_id, quality or 28)
    else:
        raise ValueError(f"Compression not supported for {original_ext}")


def _compress_image(input_path, original_ext, upload_folder, unique_id, quality):
    output_ext = original_ext if original_ext not in ('heic', 'jfif') else 'jpg'
    output_filename = f"{unique_id}_compressed.{output_ext}"
    output_path = os.path.join(upload_folder, output_filename)

    with Image.open(input_path) as img:
        # Flatten transparency for JPEG
        if output_ext in ('jpg', 'jpeg') and img.mode in ('RGBA', 'P', 'LA'):
            bg = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            if img.mode in ('RGBA', 'LA'):
                bg.paste(img, mask=img.split()[-1])
            img = bg

        fmt_map = {'jpg': 'JPEG', 'jpeg': 'JPEG', 'jfif': 'JPEG', 'tif': 'TIFF'}
        fmt = fmt_map.get(output_ext, output_ext.upper())

        if fmt == 'PNG':
            # PNG uses lossless compression level 0-9; map quality 0-100 → 9-0
            compress_level = max(0, min(9, int((100 - quality) / 11)))
            img.save(output_path, format='PNG', optimize=True, compress_level=compress_level)
        elif fmt in ('JPEG',):
            img.save(output_path, format=fmt, quality=int(quality), optimize=True)
        else:
            img.save(output_path, format=fmt)

    return output_filename, output_path


def _compress_pdf(input_path, upload_folder, unique_id, quality):
    """
    Compress PDF by re-encoding embedded images at reduced quality
    and removing redundant objects.
    """
    output_filename = f"{unique_id}_compressed.pdf"
    output_path = os.path.join(upload_folder, output_filename)

    reader = PdfReader(input_path)
    writer = PdfWriter()

    for page in reader.pages:
        writer.add_page(page)

    # Compress each page's content streams
    for page in writer.pages:
        page.compress_content_streams()

    # Re-encode images at lower quality
    img_quality = max(10, min(95, int(quality)))
    for page in writer.pages:
        if '/Resources' in page and '/XObject' in page['/Resources']:
            xobjects = page['/Resources']['/XObject']
            if hasattr(xobjects, 'items'):
                for name, obj in xobjects.items():
                    try:
                        xobj = obj.get_object()
                        if xobj.get('/Subtype') == '/Image':
                            width = int(xobj['/Width'])
                            height = int(xobj['/Height'])
                            data = xobj._data
                            if xobj.get('/ColorSpace') == '/DeviceRGB':
                                mode = 'RGB'
                            else:
                                mode = 'L'
                            try:
                                img = Image.frombytes(mode, (width, height), data)
                                buf = io.BytesIO()
                                img.save(buf, format='JPEG', quality=img_quality)
                                xobj._data = buf.getvalue()
                                xobj['/Filter'] = '/DCTDecode'
                                xobj['/Length'] = len(xobj._data)
                            except Exception:
                                pass  # skip unreadable images
                    except Exception:
                        pass

    with open(output_path, 'wb') as f:
        writer.write(f)

    return output_filename, output_path


def _compress_audio(input_path, original_ext, upload_folder, unique_id, bitrate_kbps):
    output_filename = f"{unique_id}_compressed.{original_ext}"
    output_path = os.path.join(upload_folder, output_filename)

    audio = AudioSegment.from_file(input_path)
    bitrate = f"{int(bitrate_kbps)}k"
    fmt_map = {'m4a': 'ipod'}
    fmt = fmt_map.get(original_ext, original_ext)
    audio.export(output_path, format=fmt, bitrate=bitrate)
    return output_filename, output_path


def _compress_video(input_path, original_ext, upload_folder, unique_id, crf):
    """CRF: lower = better quality / larger file. 28 is a sensible default."""
    output_filename = f"{unique_id}_compressed.{original_ext}"
    output_path = os.path.join(upload_folder, output_filename)

    try:
        quality_val = int(crf)
    except (TypeError, ValueError):
        quality_val = 28

    # If caller supplies a 1-100 quality percentage, map that into the CRF range.
    if 1 <= quality_val <= 100:
        crf_value = max(18, min(35, int(35 - (quality_val / 100) * 17)))
    else:
        crf_value = max(18, min(35, quality_val))

    video = VideoFileClip(input_path)
    # moviepy doesn't expose CRF directly — we reduce resolution as a proxy
    # Scale down to 720p if larger
    if video.h > 720:
        video = video.resize(height=720)

    video.write_videofile(output_path, codec='libx264',
                          ffmpeg_params=['-crf', str(crf_value)],
                          logger=None)
    video.close()
    return output_filename, output_path
